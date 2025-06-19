from flask import Blueprint, render_template, request, redirect, url_for, flash, session, jsonify, send_file
from database import db
from models import Employee, School, Penalty, InvestigationCommittee
from datetime import datetime
from routes.auth_routes import login_required, permission_required
from utils import log_user_activity
from sqlalchemy import or_
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import io
import os

procedures = Blueprint('procedures', __name__)

@procedures.route('/procedures')
@login_required
def procedures_list():
    """عرض قائمة الإجراءات"""
    page = request.args.get('page', 1, type=int)
    search = request.args.get('search', '')
    
    query = Penalty.query.join(Employee).join(School)
    
    if search:
        query = query.filter(
            or_(
                Employee.name.contains(search),
                Employee.ministry_number.contains(search),
                Penalty.penalty_type.contains(search),
                School.name.contains(search)
            )
        )
    
    penalties = query.order_by(Penalty.created_at.desc()).paginate(
        page=page, per_page=20, error_out=False
    )
    
    return render_template('procedures/penalties_list.html', penalties=penalties, search=search)

@procedures.route('/procedures/add_penalty', methods=['GET', 'POST'])
@login_required
def add_penalty():
    """إضافة عقوبة جديدة"""
    if request.method == 'POST':
        try:
            # الحصول على البيانات من النموذج
            ministry_number = request.form.get('ministry_number')
            penalty_type = request.form.get('penalty_type')
            school_id = request.form.get('school_id')
            letter_number = request.form.get('letter_number')
            letter_date = datetime.strptime(request.form.get('letter_date'), '%Y-%m-%d').date()
            notes = request.form.get('notes', '')
            
            # البحث عن الموظف
            employee = Employee.query.filter_by(ministry_number=ministry_number).first()
            if not employee:
                flash('الموظف غير موجود', 'error')
                return redirect(url_for('procedures.add_penalty'))
            
            # التحقق من وجود المدرسة/القسم
            school = School.query.get(school_id)
            if not school:
                flash('المدرسة/القسم غير موجود', 'error')
                return redirect(url_for('procedures.add_penalty'))
            
            # إنشاء العقوبة الجديدة
            penalty = Penalty(
                employee_id=employee.id,
                penalty_type=penalty_type,
                school_id=school_id,
                letter_number=letter_number,
                letter_date=letter_date,
                notes=notes
            )
            
            db.session.add(penalty)
            db.session.commit()
            
            # تسجيل النشاط
            log_user_activity(session.get('user_id'), 'إضافة عقوبة', f'تم إضافة عقوبة للموظف: {employee.name}')
            
            flash('تم إضافة العقوبة بنجاح', 'success')
            return redirect(url_for('procedures.procedures_list'))
            
        except Exception as e:
            db.session.rollback()
            flash(f'حدث خطأ: {str(e)}', 'error')
    
    # الحصول على قائمة المدارس والأقسام
    schools = School.query.order_by(School.name).all()
    
    return render_template('procedures/add_penalty.html', schools=schools)

@procedures.route('/procedures/edit_penalty/<int:penalty_id>', methods=['GET', 'POST'])
@login_required
def edit_penalty(penalty_id):
    """تعديل عقوبة"""
    penalty = Penalty.query.get_or_404(penalty_id)
    
    if request.method == 'POST':
        try:
            penalty.penalty_type = request.form.get('penalty_type')
            penalty.school_id = request.form.get('school_id')
            penalty.letter_number = request.form.get('letter_number')
            penalty.letter_date = datetime.strptime(request.form.get('letter_date'), '%Y-%m-%d').date()
            penalty.notes = request.form.get('notes', '')
            penalty.updated_at = datetime.utcnow()
            
            db.session.commit()
            
            # تسجيل النشاط
            log_user_activity(session.get('user_id'), 'تعديل عقوبة', f'تم تعديل عقوبة للموظف: {penalty.employee.name}')
            
            flash('تم تعديل العقوبة بنجاح', 'success')
            return redirect(url_for('procedures.procedures_list'))
            
        except Exception as e:
            db.session.rollback()
            flash(f'حدث خطأ: {str(e)}', 'error')
    
    schools = School.query.order_by(School.name).all()
    return render_template('procedures/edit_penalty.html', penalty=penalty, schools=schools)

@procedures.route('/procedures/delete_penalty/<int:penalty_id>', methods=['POST'])
@login_required
def delete_penalty(penalty_id):
    """حذف عقوبة"""
    penalty = Penalty.query.get_or_404(penalty_id)
    
    try:
        employee_name = penalty.employee.name
        db.session.delete(penalty)
        db.session.commit()
        
        # تسجيل النشاط
        log_user_activity(session.get('user_id'), 'حذف عقوبة', f'تم حذف عقوبة للموظف: {employee_name}')
        
        flash('تم حذف العقوبة بنجاح', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}', 'error')
    
    return redirect(url_for('procedures.procedures_list'))

@procedures.route('/procedures/print_penalty/<int:penalty_id>')
@login_required
def print_penalty(penalty_id):
    """طباعة العقوبة"""
    penalty = Penalty.query.get_or_404(penalty_id)
    return render_template('procedures/print_penalty.html', penalty=penalty)

@procedures.route('/procedures/export_penalty_word/<int:penalty_id>')
@login_required
def export_penalty_word(penalty_id):
    """تصدير العقوبة إلى ملف Word"""
    penalty = Penalty.query.get_or_404(penalty_id)
    
    # إنشاء مستند Word جديد
    doc = Document()
    
    # إعداد اتجاه النص من اليمين لليسار
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11.69)  # A4
        section.page_width = Inches(8.27)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
    
    # إضافة الهيدر
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_p.add_run('المملكة الأردنية الهاشمية\nمديرية التربية والتعليم لمنطقة الزرقاء الثانية')
    header_run.bold = True
    header_run.font.size = Inches(0.16)
    
    # إضافة مساحة
    doc.add_paragraph()
    
    # معلومات المستند
    doc_info_p = doc.add_paragraph()
    doc_info_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc_info_p.add_run('الرقم: ........................\nالتاريخ: ........................\nالموافق: ........................')
    
    # العنوان
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('معالي وزير التربية والتعليم المحترم')
    title_run.bold = True
    title_run.font.size = Inches(0.18)
    
    doc.add_paragraph()
    
    # الموضوع
    subject_p = doc.add_paragraph()
    subject_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subject_p.add_run('الموضوع: الإجراءات والعقوبات التأديبية')
    
    # معلومات الموظف
    employee_gender = penalty.employee.gender
    
    # معالجة خاصة للمسميات الوظيفية
    job_title = penalty.employee.job_title
    
    # قائمة المسميات الوظيفية التي تحتاج معالجة خاصة
    special_titles = {
        "مدير المدرسة": {"male": "مدير المدرسة", "female": "مديرة المدرسة"},
        "مدير": {"male": "المدير", "female": "المديرة"},
        "نائب مدير": {"male": "نائب المدير", "female": "نائبة المدير"},
        "رئيس قسم": {"male": "رئيس القسم", "female": "رئيسة القسم"},
        "مساعد مدير": {"male": "مساعد المدير", "female": "مساعدة المدير"}
    }
    
    if job_title in special_titles:
        if employee_gender in ['أنثى', 'انثى']:
            job_title_with_gender = special_titles[job_title]["female"]
        else:
            job_title_with_gender = special_titles[job_title]["male"]
    else:
        # للمسميات الوظيفية العادية
        if employee_gender in ['أنثى', 'انثى']:
            if job_title.startswith('ال'):
                job_title_with_gender = f"{job_title}ة"
            else:
                job_title_with_gender = f"ال{job_title}ة"
        else:
            if job_title.startswith('ال'):
                job_title_with_gender = job_title
            else:
                job_title_with_gender = f"ال{job_title}"
    
    employee_p = doc.add_paragraph()
    employee_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    employee_p.add_run(f'{job_title_with_gender}/{penalty.employee.name} {penalty.employee.ministry_number}')
    
    doc.add_paragraph()
    
    # المحتوى
    content_p = doc.add_paragraph()
    content_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    content_p.add_run('السلام عليكم ورحمة الله وبركاته')
    
    doc.add_paragraph()
    
    # النص الرئيسي
    main_text = f"أرفق طيه عقوبة {penalty.penalty_type} الصادرة بحق "
    main_text += "المذكورة اعلاه" if employee_gender in ['أنثى', 'انثى'] else "المذكور اعلاه"
    main_text += f" بناءً على كتاب "
    main_text += "مديرة مدرستها" if employee_gender in ['أنثى', 'انثى'] else "مدير مدرسته"
    main_text += f" رقم {penalty.letter_number} تاريخ {penalty.letter_date.strftime('%Y/%m/%d')}م."
    
    main_p = doc.add_paragraph()
    main_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    main_p.add_run(main_text)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # الاحترام
    respect_p = doc.add_paragraph()
    respect_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    respect_run = respect_p.add_run('واقبلوا الاحترام')
    respect_run.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # التوقيع
    signature_p = doc.add_paragraph()
    signature_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    signature_run = signature_p.add_run('مدير التربية والتعليم')
    signature_run.bold = True
    
    # النسخ
    doc.add_paragraph()
    doc.add_paragraph()
    
    copies_text = """نسخة/ معالي رئيس ديوان المحاسبة المحترم
نسخة/عطوفة رئيس مجلس هيئة الخدمة والادارة العامة المحترم
نسخة / رئيس قسم التحقيقات والاجراءات /الوزارة
مع نسخة من اوراق التحقيق
نسخة مدير الشؤون الإدارية والمالية
نسخة / رئيس قسم شؤون الموظفين
نسخة / ملف الشخصي / مع المرفقات"""
    
    if penalty.employee.is_directorate_employee:
        copies_text += f"\nنسخة / لقسم {penalty.employee.department}"
    else:
        copies_text += f"\nنسخة / لمدرسة {penalty.school.name}"
    
    copies_text += "\nنسخة / الديوان"
    
    copies_p = doc.add_paragraph()
    copies_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    copies_p.add_run(copies_text)
    
    # حفظ المستند في الذاكرة
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # إنشاء اسم الملف
    filename = f"عقوبة_{penalty.employee.name}_{penalty.penalty_type}_{penalty.letter_date.strftime('%Y-%m-%d')}.docx"
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

# راوتات لجان التحقيق
@procedures.route('/procedures/investigation_committees')
@login_required
def investigation_committees_list():
    """عرض قائمة لجان التحقيق"""
    page = request.args.get('page', 1, type=int)
    search = request.args.get('search', '')
    
    query = InvestigationCommittee.query.join(School)
    
    if search:
        query = query.filter(
            or_(
                School.name.contains(search),
                InvestigationCommittee.complaint_content.contains(search),
                InvestigationCommittee.letter_number.contains(search)
            )
        )
    
    committees = query.order_by(InvestigationCommittee.created_at.desc()).paginate(
        page=page, per_page=20, error_out=False
    )
    
    return render_template('procedures/investigation_committees_list.html', committees=committees, search=search)

@procedures.route('/procedures/add_investigation_committee', methods=['GET', 'POST'])
@login_required
def add_investigation_committee():
    """إضافة لجنة تحقيق جديدة"""
    if request.method == 'POST':
        try:
            # الحصول على البيانات من النموذج
            school_id = request.form.get('school_id')
            committee_head_id = request.form.get('committee_head_id')
            first_member_id = request.form.get('first_member_id')
            second_member_id = request.form.get('second_member_id')
            complaint_content = request.form.get('complaint_content')
            committee_date = datetime.strptime(request.form.get('committee_date'), '%Y-%m-%d').date()
            letter_number = request.form.get('letter_number', '')
            letter_date_str = request.form.get('letter_date')
            letter_date = datetime.strptime(letter_date_str, '%Y-%m-%d').date() if letter_date_str else None
            notes = request.form.get('notes', '')
            
            # التحقق من أن الأعضاء مختلفون
            if committee_head_id == first_member_id or committee_head_id == second_member_id or first_member_id == second_member_id:
                flash('يجب أن يكون أعضاء اللجنة مختلفين', 'error')
                return redirect(url_for('procedures.add_investigation_committee'))
            
            # إنشاء لجنة التحقيق الجديدة
            committee = InvestigationCommittee(
                school_id=school_id,
                committee_head_id=committee_head_id,
                first_member_id=first_member_id,
                second_member_id=second_member_id,
                complaint_content=complaint_content,
                committee_date=committee_date,
                letter_number=letter_number,
                letter_date=letter_date,
                notes=notes
            )
            
            db.session.add(committee)
            db.session.commit()
            
            # تسجيل النشاط
            log_user_activity(session.get('user_id'), 'إضافة لجنة تحقيق', f'تم تشكيل لجنة تحقيق جديدة في {committee.school.name}')
            
            flash('تم تشكيل لجنة التحقيق بنجاح', 'success')
            return redirect(url_for('procedures.investigation_committees_list'))
            
        except Exception as e:
            db.session.rollback()
            flash(f'حدث خطأ أثناء تشكيل لجنة التحقيق: {str(e)}', 'error')
    
    # الحصول على المدارس والموظفين للنموذج
    schools = School.query.order_by(School.name).all()
    # جلب موظفي المديرية فقط
    employees = Employee.query.filter_by(is_directorate_employee=True).order_by(Employee.name).all()
    
    return render_template('procedures/add_investigation_committee.html', schools=schools, employees=employees)

@procedures.route('/procedures/edit_investigation_committee/<int:committee_id>', methods=['GET', 'POST'])
@login_required
def edit_investigation_committee(committee_id):
    """تعديل لجنة تحقيق"""
    committee = InvestigationCommittee.query.get_or_404(committee_id)
    
    if request.method == 'POST':
        try:
            committee.school_id = request.form.get('school_id')
            committee.committee_head_id = request.form.get('committee_head_id')
            committee.first_member_id = request.form.get('first_member_id')
            committee.second_member_id = request.form.get('second_member_id')
            committee.complaint_content = request.form.get('complaint_content')
            committee.committee_date = datetime.strptime(request.form.get('committee_date'), '%Y-%m-%d').date()
            committee.letter_number = request.form.get('letter_number', '')
            letter_date_str = request.form.get('letter_date')
            committee.letter_date = datetime.strptime(letter_date_str, '%Y-%m-%d').date() if letter_date_str else None
            committee.status = request.form.get('status')
            committee.notes = request.form.get('notes', '')
            committee.updated_at = datetime.utcnow()
            
            # التحقق من أن الأعضاء مختلفون
            if (committee.committee_head_id == committee.first_member_id or 
                committee.committee_head_id == committee.second_member_id or 
                committee.first_member_id == committee.second_member_id):
                flash('يجب أن يكون أعضاء اللجنة مختلفين', 'error')
                return redirect(url_for('procedures.edit_investigation_committee', committee_id=committee_id))
            
            db.session.commit()
            
            # تسجيل النشاط
            log_user_activity(session.get('user_id'), 'تعديل لجنة تحقيق', f'تم تعديل لجنة التحقيق في {committee.school.name}')
            
            flash('تم تعديل لجنة التحقيق بنجاح', 'success')
            return redirect(url_for('procedures.investigation_committees_list'))
            
        except Exception as e:
            db.session.rollback()
            flash(f'حدث خطأ أثناء تعديل لجنة التحقيق: {str(e)}', 'error')
    
    # الحصول على المدارس والموظفين للنموذج
    schools = School.query.order_by(School.name).all()
    # جلب موظفي المديرية فقط
    employees = Employee.query.filter_by(is_directorate_employee=True).order_by(Employee.name).all()
    
    return render_template('procedures/edit_investigation_committee.html', committee=committee, schools=schools, employees=employees)

@procedures.route('/procedures/delete_investigation_committee/<int:committee_id>', methods=['POST'])
@login_required
def delete_investigation_committee(committee_id):
    """حذف لجنة تحقيق"""
    committee = InvestigationCommittee.query.get_or_404(committee_id)
    
    try:
        school_name = committee.school.name
        db.session.delete(committee)
        db.session.commit()
        
        # تسجيل النشاط
        log_user_activity(session.get('user_id'), 'حذف لجنة تحقيق', f'تم حذف لجنة التحقيق في {school_name}')
        
        flash('تم حذف لجنة التحقيق بنجاح', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ أثناء حذف لجنة التحقيق: {str(e)}', 'error')
    
    return redirect(url_for('procedures.investigation_committees_list'))

@procedures.route('/procedures/print_investigation_committee/<int:committee_id>')
@login_required
def print_investigation_committee(committee_id):
    """طباعة لجنة التحقيق"""
    committee = InvestigationCommittee.query.get_or_404(committee_id)
    return render_template('procedures/print_investigation_committee.html', committee=committee)

# API للبحث عن موظفي المديرية
@procedures.route('/api/directorate_employee_search')
@login_required
def api_directorate_employee_search():
    try:
        search_term = request.args.get('search_term', '').strip()
        
        if not search_term:
            return jsonify({
                'success': False,
                'message': 'يرجى إدخال كلمة البحث'
            })
        
        # البحث عن موظفي المديرية بالاسم أو الرقم الوزاري
        query = Employee.query.filter_by(is_directorate_employee=True)
        
        query = query.filter(
            db.or_(
                Employee.name.ilike(f'%{search_term}%'),
                Employee.ministry_number.ilike(f'%{search_term}%')
            )
        )
        
        employees = query.limit(10).all()
        
        if employees:
            # تحويل النتائج إلى قائمة
            employees_list = []
            for employee in employees:
                employee_data = {
                    'id': employee.id,
                    'name': employee.name,
                    'ministry_number': employee.ministry_number,
                    'job_title': employee.job_title
                }
                employees_list.append(employee_data)
            
            return jsonify({
                'success': True,
                'employees': employees_list
            })
        else:
            return jsonify({
                'success': False,
                'message': 'لم يتم العثور على موظفين مطابقين لكلمة البحث'
            })
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': f'حدث خطأ أثناء البحث: {str(e)}'
        })

@procedures.route('/api/search_employee')
@login_required
def search_employee():
    """البحث عن موظف بالرقم الوزاري - API"""
    ministry_number = request.args.get('ministry_number', '')
    
    if ministry_number:
        employee = Employee.query.filter_by(ministry_number=ministry_number).first()
        if employee:
            return {
                'success': True,
                'employee': {
                    'id': employee.id,
                    'name': employee.name,
                    'ministry_number': employee.ministry_number,
                    'school_id': employee.school_id if employee.school_id else None,
                    'school_name': employee.school.name if employee.school else ''
                }
            }
    
    return {'success': False, 'message': 'الموظف غير موجود'}

@procedures.route('/procedures/export_penalty_word_formatted/<int:penalty_id>')
@login_required
def export_penalty_word_formatted(penalty_id):
    """تصدير العقوبة إلى ملف Word بنفس تنسيق الصفحة المطبوعة"""
    penalty = Penalty.query.get_or_404(penalty_id)
    
    # إنشاء مستند Word جديد
    doc = Document()
    
    # إعداد اتجاه النص من اليمين لليسار
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11.69)  # A4
        section.page_width = Inches(8.27)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        
        # إعداد اتجاه النص RTL
        sectPr = section._sectPr
        bidi = OxmlElement('w:bidi')
        sectPr.append(bidi)
    
    # إضافة الهيدر مع الشعار
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # النص العلوي
    header_run1 = header_p.add_run('المملكة الأردنية الهاشمية')
    header_run1.bold = True
    header_run1.font.size = Inches(0.14)
    header_p.add_run('\n')
    
    # مساحة للشعار (يمكن إضافة الشعار هنا إذا كان متاحاً)
    header_p.add_run('\n')
    
    # النص السفلي
    header_run2 = header_p.add_run('مديرية التربية والتعليم لمنطقة الزرقاء الثانية')
    header_run2.bold = True
    header_run2.font.size = Inches(0.14)
    
    # إضافة مساحة
    doc.add_paragraph()
    
    # معلومات المستند في الأعلى يميناً
    doc_info_p = doc.add_paragraph()
    doc_info_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc_info_run = doc_info_p.add_run('الرقم    :...........................\nالتاريخ  :...........................\nالموافق  :...........................')
    doc_info_run.font.size = Inches(0.10)
    
    # إضافة مساحة
    doc.add_paragraph()
    
    # العنوان
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('معالي وزير التربية والتعليم المحترم')
    title_run.bold = True
    title_run.font.size = Inches(0.18)
    
    # إضافة مساحة
    doc.add_paragraph()
    doc.add_paragraph()
    
    # الموضوع
    subject_p = doc.add_paragraph()
    subject_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subject_run = subject_p.add_run('الموضوع: الإجراءات والعقوبات التأديبية')
    subject_run.bold = True
    subject_run.font.size = Inches(0.16)
    
    # معلومات الموظف مع معالجة المسمى الوظيفي
    employee_gender = penalty.employee.gender
    job_title = penalty.employee.job_title
    
    # معالجة المسمى الوظيفي بنفس طريقة الصفحة المطبوعة
    if job_title == "مدير المدرسة":
        if employee_gender in ['أنثى', 'انثى']:
            formatted_job_title = "مديرة المدرسة"
        else:
            formatted_job_title = "مدير المدرسة"
    elif job_title == "مساعد مدير المدرسة":
        if employee_gender in ['أنثى', 'انثى']:
            formatted_job_title = "مساعدة مديرة المدرسة"
        else:
            formatted_job_title = "مساعد مدير المدرسة"
    elif job_title == "امين اللوازم المدرسية":
        if employee_gender in ['أنثى', 'انثى']:
            formatted_job_title = "امينة اللوازم المدرسية"
        else:
            formatted_job_title = "امين اللوازم المدرسية"
    elif job_title == "امين المكتبة":
        if employee_gender in ['أنثى', 'انثى']:
            formatted_job_title = "امينة المكتبة"
        else:
            formatted_job_title = "امين المكتبة"
    elif job_title == "قيم مختبر حاسوب":
        if employee_gender in ['أنثى', 'انثى']:
            formatted_job_title = "قيمة مختبر الحاسوب"
        else:
            formatted_job_title = "قيم مختبر حاسوب"
    elif job_title == "معلم":
        if employee_gender in ['أنثى', 'انثى']:
            formatted_job_title = "المعلمة"
        else:
            formatted_job_title = "المعلم"
    elif job_title == "مبرمج مساعد":
        if employee_gender in ['أنثى', 'انثى']:
            formatted_job_title = "المبرمجة المساعدة"
        else:
            formatted_job_title = "المبرمج المساعد"
    elif job_title == "مدير":
        if employee_gender in ['أنثى', 'انثى']:
            formatted_job_title = "المديرة"
        else:
            formatted_job_title = "المدير"
    elif job_title == "نائب مدير":
        if employee_gender in ['أنثى', 'انثى']:
            formatted_job_title = "نائبة المدير"
        else:
            formatted_job_title = "نائب المدير"
    elif job_title == "رئيس قسم":
        if employee_gender in ['أنثى', 'انثى']:
            formatted_job_title = "رئيسة القسم"
        else:
            formatted_job_title = "رئيس القسم"
    else:
        # للمسميات الوظيفية العادية
        if employee_gender in ['أنثى', 'انثى']:
            if job_title.startswith('ال'):
                formatted_job_title = f"{job_title}ة"
            else:
                formatted_job_title = f"ال{job_title}ة"
        else:
            if job_title.startswith('ال'):
                formatted_job_title = job_title
            else:
                formatted_job_title = f"ال{job_title}"
    
    employee_p = doc.add_paragraph()
    employee_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    employee_run = employee_p.add_run(f'{formatted_job_title}/{penalty.employee.name} ({penalty.employee.ministry_number})')
    employee_run.font.size = Inches(0.16)
    
    # إضافة مساحة
    doc.add_paragraph()
    
    # المحتوى
    content_p = doc.add_paragraph()
    content_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    content_run = content_p.add_run('السلام عليكم ورحمة الله وبركاته')
    content_run.font.size = Inches(0.16)
    
    # إضافة مساحة
    doc.add_paragraph()
    
    # النص الرئيسي
    main_text = f"أرفق طيه عقوبة {penalty.penalty_type} الصادرة بحق "
    main_text += "المذكورة اعلاه" if employee_gender in ['أنثى', 'انثى'] else "المذكور اعلاه"
    main_text += f" بناءً على كتاب "
    main_text += "مديرة مدرستها" if employee_gender in ['أنثى', 'انثى'] else "مدير مدرسته"
    main_text += f" رقم ({penalty.letter_number}) تاريخ {penalty.letter_date.strftime('%Y/%m/%d')}م."
    
    main_p = doc.add_paragraph()
    main_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    main_run = main_p.add_run(main_text)
    main_run.font.size = Inches(0.16)
    
    # إضافة مساحة
    doc.add_paragraph()
    doc.add_paragraph()
    
    # الاحترام
    respect_p = doc.add_paragraph()
    respect_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    respect_run = respect_p.add_run('واقبلوا الاحترام')
    respect_run.bold = True
    respect_run.font.size = Inches(0.16)
    
    # إضافة مساحة
    doc.add_paragraph()
    doc.add_paragraph()
    
    # التوقيع
    signature_p = doc.add_paragraph()
    signature_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    signature_run = signature_p.add_run('مدير التربية والتعليم')
    signature_run.bold = True
    signature_run.font.size = Inches(0.16)
    
    # إضافة مساحة كبيرة قبل النسخ
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # النسخ
    copies_text = """نسخة/ معالي رئيس ديوان المحاسبة المحترم
نسخة/عطوفة رئيس مجلس هيئة الخدمة والادارة العامة المحترم
نسخة / رئيس قسم التحقيقات والاجراءات /الوزارة
مع نسخة من اوراق التحقيق
نسخة مدير الشؤون الإدارية والمالية
نسخة / رئيس قسم شؤون الموظفين
نسخة / ملف الشخصي / مع المرفقات"""
    
    if penalty.employee.is_directorate_employee:
        copies_text += f"\nنسخة / لقسم {penalty.employee.department}"
    else:
        copies_text += f"\nنسخة / لمدرسة {penalty.school.name}"
    
    copies_text += "\nنسخة / الديوان"
    
    copies_p = doc.add_paragraph()
    copies_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    copies_run = copies_p.add_run(copies_text)
    copies_run.bold = True
    copies_run.font.size = Inches(0.12)
    
    # حفظ المستند في الذاكرة
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # إنشاء اسم الملف
    filename = f"عقوبة_مطبوعة_{penalty.employee.name}_{penalty.penalty_type}_{penalty.letter_date.strftime('%Y-%m-%d')}.docx"
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )